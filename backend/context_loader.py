"""
Context loader - auto-load all files from context directory,
extract text, merge into experiment context string.
Supports: .md .txt .pdf .docx
Auto dedup + filter personal info (phone/email)
"""
import os
import re
from typing import Optional

CONTEXT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "上下文")
os.makedirs(CONTEXT_DIR, exist_ok=True)

# 个人信息过滤正则
_PHONE_RE = re.compile(r'1[3-9]\d{9}')
_EMAIL_RE = re.compile(r'[\w.\-]+@[\w.\-]+')
_QQ_RE = re.compile(r'\d{5,12}@qq\.com')
_BIRTH_RE = re.compile(r'\d{4}/\d{2}')
_PURE_TABLE_ROW = re.compile(r'^[|  \t\w\u4e00-\u9fff·]+[|][|  \t\w\u4e00-\u9fff·]+$')


def _clean_text(text: str) -> str:
    """去重 + 过滤个人信息 + 表格行拆分去重"""
    lines = text.split("\n")
    cleaned = []
    seen = set()
    prev = ""

    for line in lines:
        stripped = line.strip()

        if not stripped:
            cleaned.append(line)
            continue

        # 含 | 分隔符的表格行：拆分成独立子句，去重后保留
        if "|" in stripped:
            parts = [p.strip() for p in stripped.split("|")]
            unique_parts = []
            seen_parts = set()
            for p in parts:
                if not p or len(p) < 2:
                    continue
                # 跳过含个人信息的部分
                if _PHONE_RE.search(p) or _EMAIL_RE.search(p):
                    continue
                if "通讯地址" in p or "联系电话" in p or "所在单位" in p:
                    continue
                if "申报者情况" in p or "合作者情况" in p:
                    continue
                if re.match(r'^[男女]$|^\d{4}/\d{2}$|^\d{4}$|^\d{4}级$', p):
                    continue
                if re.match(r'^(本科|硕士|博士|专科|中级|高级)$', p):
                    continue
                # 去重
                key = p[:60]
                if key not in seen_parts and key not in seen:
                    unique_parts.append(p)
                    seen_parts.add(key)

            for p in unique_parts:
                norm = p[:80]
                if norm in seen:
                    continue
                if p == prev:
                    continue
                seen.add(norm)
                prev = p
                cleaned.append(p)
            continue

        # 跳过含手机号、邮箱的行
        if _PHONE_RE.search(stripped) or _EMAIL_RE.search(stripped):
            continue

        # 跳过个人信息行
        if "通讯地址" in stripped or "联系电话" in stripped or "所在单位" in stripped:
            continue
        if re.match(r'^[男女]\s+\d{4}', stripped):
            continue
        if re.match(r'^[\u4e00-\u9fff·]{2,4}\s+[\u4e00-\u9fff·]{2,4}\s+[男女]', stripped):
            continue
        if "申报者情况" in stripped or "合作者情况" in stripped or "指导教师" in stripped:
            continue

        # 去重
        if stripped == prev:
            continue
        prev = stripped

        norm = stripped[:80]
        if norm in seen:
            continue
        seen.add(norm)

        cleaned.append(line)

    result = "\n".join(cleaned).strip()
    result = re.sub(r'\n{4,}', '\n\n\n', result)
    return result


def _read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        text = f.read()
    # .md 文件不需要去重清洗（用户手写的），只清洗 .docx 提取的结果
    return text


def _read_pdf(path: str) -> str:
    # 先试 pdfplumber
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
        result = "\n\n".join(pages)
        if result.strip():
            return _clean_text(result)
    except Exception:
        pass

    # fallback: PyPDF2
    try:
        from PyPDF2 import PdfReader
        pages = []
        reader = PdfReader(path)
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                pages.append(text.strip())
        result = "\n\n".join(pages)
        if result.strip():
            return _clean_text(result)
    except Exception:
        pass

    return ""


def _read_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    # also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    raw = "\n".join(paragraphs)
    return _clean_text(raw)


READERS = {
    ".md": _read_txt,
    ".txt": _read_txt,
    ".pdf": _read_pdf,
    ".docx": _read_docx,
}


def load_context() -> tuple[str, list[str], list[str]]:
    """
    扫描 上下文 目录，读取所有支持的文件，合并为上下文字符串。
    返回: (上下文文本, 已加载文件列表, 空/失败文件列表)
    """
    if not os.path.isdir(CONTEXT_DIR):
        return "", [], []

    parts = []
    loaded = []
    skipped = []
    files = sorted(os.listdir(CONTEXT_DIR))

    for filename in files:
        ext = os.path.splitext(filename)[1].lower()
        if ext not in READERS:
            continue

        filepath = os.path.join(CONTEXT_DIR, filename)
        try:
            text = READERS[ext](filepath)
            if text.strip():
                label = os.path.splitext(filename)[0]
                parts.append(f"## {label}\n\n{text.strip()}")
                loaded.append(filename)
                print(f"[context] Loaded: {filename} ({len(text)} chars)")
            else:
                skipped.append(filename)
                print(f"[context] Skipped (no extractable text): {filename}")
        except Exception as e:
            skipped.append(filename)
            print(f"[context] Failed: {filename} - {e}")

    if skipped:
        summary = "\n\n> 以下文件未能提取文本（可能是扫描件/图片PDF）：\n"
        for f in skipped:
            summary += f"> - {f}\n"
        parts.append(summary)

    return "\n\n---\n\n".join(parts), loaded, skipped


def get_context(force_reload: bool = False) -> str:
    ctx, loaded, skipped = load_context()
    return ctx if ctx else ""


def get_notes() -> str:
    """
    返回实验说明页的核心内容。
    优先读取 实验说明.md，不存在则返回空。
    """
    if not os.path.isdir(CONTEXT_DIR):
        return ""
    for filename in sorted(os.listdir(CONTEXT_DIR)):
        if filename == "实验说明.md":
            filepath = os.path.join(CONTEXT_DIR, filename)
            try:
                return _read_txt(filepath).strip()
            except Exception:
                pass
    return ""


def reload_context() -> str:
    return get_context(force_reload=True)


def get_context_file_count() -> dict:
    if not os.path.isdir(CONTEXT_DIR):
        return {"total": 0, "loaded": 0, "skipped": 0}
    _, loaded, skipped = load_context()
    return {"total": len(loaded) + len(skipped), "loaded": len(loaded), "skipped": len(skipped)}


if __name__ == "__main__":
    ctx, ld, sk = load_context()
    print(f"\nDone. {len(ctx)} chars, loaded={ld}, skipped={sk}")
